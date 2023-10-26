from flask import Blueprint,redirect, url_for, flash,render_template, request, session, render_template_string, send_file
from flask_login import login_required, current_user
from .models import BusinessUpdates
from .models import Users, Portfolios , Projects
import smtplib
from email.message import EmailMessage
from docx import Document
import io
import re
import pandas as pd
import tempfile
import os
from docx.shared import RGBColor

admin_view = Blueprint('admin_view', __name__)

EMAIL_ADDRESS = "reportgenrator@gmail.com"
EMAIL_PASSWORD = os.getenv('EMAIL_PASSWORD')

@admin_view.route('/admin_landing')
@login_required
def admin_landing():
    if not current_user.is_authenticated:
        return redirect(url_for('auth.login'))  # Redirect to the login page
    return render_template('admin_landing.html')
@login_required
@admin_view.route('/excel', methods=['POST'])
def excel():
    todate = request.form['toDate']
    session['todate'] = todate
    fromdate = request.form['fromDate']
    session['fromdate'] = fromdate
    portfolio_name = request.form['project']
    session['project'] = portfolio_name

    session['service'] = request.form['services']

    query = ''
    df = ''

    if portfolio_name == 'all':  # Check against the name, not the entire object
        updates = BusinessUpdates.query.filter(
            BusinessUpdates.date.between(fromdate, todate),
            BusinessUpdates.service.like(session['service'])
        ).all()
    else:
        portfolio = Portfolios.query.filter_by(name=portfolio_name).first()
        portfolio_id = portfolio.id
        session['portfolio_id'] = portfolio_id  # Store only the portfolio_id in the session
        
        updates = BusinessUpdates.query.filter(
            BusinessUpdates.date.between(fromdate, todate),
            BusinessUpdates.portfolio_id == session['portfolio_id'],
            BusinessUpdates.portfolio_id == session['portfolio_id'],
            BusinessUpdates.service.like(session['service'])
        ).all()


    
    updates_data = []
    for update in updates:
        update_data = {
            "DATE": update.date,
            "USERNAME": update.user.username,
            "USER_INPUT": update.user_input,
            "USER_OUTPUT": update.user_output,
            "SERVICE": update.service,
            "PORTFOLIO": (Portfolios.query.filter_by(id = update.portfolio_id).first()).name,
            "PROJECT":(Projects.query.filter_by(id = update.project_id).first().name),
            "PROGRESS": update.progress,
            "TEAMMATES": update.teammates,
            "AI-BUSINESS-INPUT": update.ai_input,
            "AI-BUSINESS-OUTPUT": update.ai_output,
            "BUSINESS-UPDATE": update.business_update
        }
        updates_data.append(update_data)

    # Convert the list of dictionaries to a pandas DataFrame
    df = pd.DataFrame(updates_data)

    table_html = df.to_html(classes='table table-bordered table-striped', index=False)

    return render_template('report.html', table_html=table_html)


@login_required
@admin_view.route('/portfolio_details', methods=['GET', 'POST'])
def portfolio_details():
    fromdate = session.get('fromdate', '')
    todate = session.get('todate', '')
    service = session.get('service', '')
    portfolio_name = session.get('project', '')
    
    if portfolio_name == 'all':  # Check against the name, not the entire object
        updates = BusinessUpdates.query.filter(
            BusinessUpdates.date.between(fromdate, todate),
            BusinessUpdates.service.like(session['service'])
        ).all()
    else:
        portfolio = Portfolios.query.filter_by(name=portfolio_name).first()
        portfolio_id = portfolio.id
        session['portfolio_id'] = portfolio_id  # Store only the portfolio_id in the session
        
        updates = BusinessUpdates.query.filter(
            BusinessUpdates.date.between(fromdate, todate),
            BusinessUpdates.portfolio_id == session['portfolio_id'],
            BusinessUpdates.service.like(session['service'])
        ).all()


    updates_data = []
    for update in updates:
        update_data = {
            "Date": update.date,
            "Service": update.service,
            "AI_Input": update.ai_input,
            "Business_Update": update.business_update,
            "Portfolio": (Portfolios.query.filter_by(id = update.portfolio_id).first()).name,
            "Project":(Projects.query.filter_by(id = update.project_id).first().name)
            # Add other columns as needed
        }
        updates_data.append(update_data)

    # Convert the list of dictionaries to a pandas DataFrame
    df = pd.DataFrame(updates_data)
    print(df)
    if df.empty:
        flash("No data available", category="error")
        return redirect(url_for("admin_view.admin_landing"))

    portfolio_details = ""
    

    list1 = []
    for name in df["Portfolio"]:
        if name in list1:
            continue
        else:
            list1.append(name)

    finalstr = ""
    for x in list1:
        finalstr = finalstr + "\n""\n" + x+ "\n"
        for index, row in df.iterrows():
            if x == row["Portfolio"]:
                date_ = str(row['Date'])
                finalstr += f"""
        {row['Project']} - ({date_[0:10]})
        {row['AI_Input']} - {row['Business_Update']}
    """
    

    return render_template_string(render_template('portfolio_details.html', portfolio_details=finalstr))


@login_required
@admin_view.route('/updated_portfolio_details', methods=['GET', 'POST'])
def update_portfolio_details():
    portfolio_details = request.form.get('portfolio-textarea')
    session["portfolio-textarea"] = portfolio_details

    return render_template_string(render_template('updated_portfolio_details.html', portfolio_details=portfolio_details))

@login_required
@admin_view.route('/download_portfolio_docx', methods=['POST'])
def download_portfolio_docx():
    fromdate = session.get('fromdate', '')
    todate = session.get('todate', '')
    service = session.get('service', '')
    portfolio_name = session.get('project', '')
    portfolio_details = session.get('portfolio-textarea' , '')

    if portfolio_name == 'all':  # Check against the name, not the entire object
        updates = BusinessUpdates.query.filter(
            BusinessUpdates.date.between(fromdate, todate),
            BusinessUpdates.service.like(session['service'])
        ).all()
    else:
        portfolio = Portfolios.query.filter_by(name=portfolio_name).first()
        portfolio_id = portfolio.id
        session['portfolio_id'] = portfolio_id  # Store only the portfolio_id in the session
        
        updates = BusinessUpdates.query.filter(
            BusinessUpdates.date.between(fromdate, todate),
            BusinessUpdates.portfolio_id == session['portfolio_id'],
            BusinessUpdates.service.like(session['service'])
        ).all()


    updates_data = []
    for update in updates:
        update_data = {
            "Date": update.date,
            "Service": update.service,
            "AI_Input": update.ai_input,
            "Business_Update": update.business_update,
            "Portfolio": (Portfolios.query.filter_by(id = update.portfolio_id).first()).name,
            "Project":(Projects.query.filter_by(id = update.project_id).first().name)
            # Add other columns as needed
        }
        updates_data.append(update_data)

    # Convert the list of dictionaries to a pandas DataFrame
    df = pd.DataFrame(updates_data)
    print(df)
    temp_doc = Document()
    temp_doc.add_heading('Project Updates', level=0)
    text = portfolio_details
    # list1 = []

    # for name in df["Portfolio"]:
    #     if name in list1:
    #         continue
    #     else:
    #         list1.append(name)
            
    # finalstr = ""
    # for x in list1:
    #     finalstr = finalstr +"\n""\n"+ x
    #     temp_doc.add_heading(x, level=2)
    #     for index, row in df.iterrows():
    #         if x == row["Portfolio"]:
    #             detail = f"""
    #     {row['AI_Input']} - {row['Business_Update']}
    # """         
    #             da_te = str(row['Date'])
    #             project = f"{row['Project']} - ({da_te[0:10]})"
    #             temp_doc.add_heading(project, level=4 )
    #             temp_doc.add_paragraph(detail)

    sections = re.split(r'(\w+\s*\w*)\n', text)
    headings = re.findall(r'(\w+\s*\w*)\n', text)
    # Extract and print only the content
    for heading, content in zip(headings, sections[2::2]):                  
        temp_doc.add_heading(f'{heading.strip()}'+ '\n', level=2)
        temp_doc.add_paragraph(content)

            
    temp_doc_file = io.BytesIO()
    temp_doc.save(temp_doc_file)
    temp_doc_file.seek(0)

    return send_file(
        temp_doc_file,
        as_attachment=True,
        download_name='portfolio_details.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )



@login_required
@admin_view.route('/download_xlsx', methods=['GET', 'POST'])
def download_xlsx():
    fromdate = session.get('fromdate', '')
    todate = session.get('todate', '')
    service = session.get('service', '')
    portfolio_name = session.get('project', '')

    df = ''
    if portfolio_name == 'all':  # Check against the name, not the entire object
        updates = BusinessUpdates.query.filter(
            BusinessUpdates.date.between(fromdate, todate),
            BusinessUpdates.service.like(session['service'])
        ).all()
    else:
        portfolio = Portfolios.query.filter_by(name=portfolio_name).first()
        portfolio_id = portfolio.id
        session['portfolio_id'] = portfolio_id  # Store only the portfolio_id in the session
        
        updates = BusinessUpdates.query.filter(
            BusinessUpdates.date.between(fromdate, todate),
            BusinessUpdates.portfolio_id == session['portfolio_id'],
            BusinessUpdates.service.like(session['service'])
        ).all()


    
    updates_data = []
    for update in updates:
        update_data = {
            "DATE": update.date,
            "USERNAME": update.user.username,
            "USER_INPUT": update.user_input,
            "USER_OUTPUT": update.user_output,
            "SERVICE": update.service,
            "PORTFOLIO": (Portfolios.query.filter_by(id = update.portfolio_id).first()).name,
            "PROJECT": (Projects.query.filter_by(id = update.project_id).first().name),
            "PROGRESS": update.progress,
            "TEAMMATES": update.teammates,
            "AI-BUSINESS-INPUT": update.ai_input,
            "AI-BUSINESS-OUTPUT": update.ai_output,
            "BUSINESS-UPDATE": update.business_update
        }
        updates_data.append(update_data)
    df = pd.DataFrame(updates_data)

    temp_dir = tempfile.mkdtemp()

    xlsx_file_path = os.path.join(temp_dir, 'data.xlsx')
    with pd.ExcelWriter(xlsx_file_path, engine='xlsxwriter') as writer:
        df.to_excel(writer, sheet_name='Sheet1', index=False)

    return send_file(xlsx_file_path, as_attachment=True, download_name='data.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

@login_required
@admin_view.route("/send", methods=["GET","POST"])
def send():
    recevier_list=['terabaapm1231@gmail.com']
    # Get the portfolio_details from the session and attach it as a DOCX file
    fromdate = session.get('fromdate', '')
    todate = session.get('todate', '')
    service = session.get('service', '')
    portfolio_name = session.get('project', '')

    if portfolio_name == 'all':  # Check against the name, not the entire object
        updates = BusinessUpdates.query.filter(
            BusinessUpdates.date.between(fromdate, todate),
            BusinessUpdates.service.like(session['service'])
        ).all()
    else:
        portfolio = Portfolios.query.filter_by(name=portfolio_name).first()
        portfolio_id = portfolio.id
        session['portfolio_id'] = portfolio_id  # Store only the portfolio_id in the session
        
        updates = BusinessUpdates.query.filter(
            BusinessUpdates.date.between(fromdate, todate),
            BusinessUpdates.portfolio_id == session['portfolio_id'],
            BusinessUpdates.service.like(session['service'])
        ).all()


    updates_data = []
    for update in updates:
        update_data = {
            "Date": update.date,
            "Service": update.service,
            "AI_Input": update.ai_input,
            "Business_Update": update.business_update,
            "Portfolio": (Portfolios.query.filter_by(id = update.portfolio_id).first()).name,
            "Project":(Projects.query.filter_by(id = update.project_id).first().name)
            # Add other columns as needed
        }
        updates_data.append(update_data)

    # Convert the list of dictionaries to a pandas DataFrame
    df = pd.DataFrame(updates_data)
    doc = Document()
    doc.add_heading('Project Updates', level=0)

    list1 = []

    for name in df["Portfolio"]:
        if name in list1:
            continue
        else:
            list1.append(name)
            
    finalstr = ""
    for x in list1:
        finalstr = finalstr +"\n""\n"+ x
        doc.add_heading(x, level=2)
        for index, row in df.iterrows():
            if x == row["Portfolio"]:
                detail = f"""
        {row['AI_Input']} - {row['Business_Update']}
    """         
                da_te = str(row['Date'])
                project = f"{row['Project']} - ({da_te[0:10]})"
                doc.add_heading(project, level=4 )
                doc.add_paragraph(detail)

    temp_docx_file = io.BytesIO()
    doc.save(temp_docx_file)
    temp_docx_file.seek(0)
    
    for item in recevier_list:
        msg = EmailMessage()  # Create a new EmailMessage instance for each recipient
        msg['Subject'] = 'This weeks Report'
        msg['From'] = EMAIL_ADDRESS
        msg['To'] = item

        msg.set_content('Hello, find this week\'s business update attached below.')
        
        # Attach the DOCX content to the email
        msg.add_attachment(
            temp_docx_file.read(),
            maintype='application',
            subtype='vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename='portfolio_details.docx'
        )

        # Send the email
        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
            smtp.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
            smtp.send_message(msg)

    flash(f"Updates sent successfully",category='success')

    return redirect(url_for("auth.logout"))